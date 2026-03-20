import os
from pathlib import Path
from docx import Document
from pypdf import PdfReader


def read_all_docx(input_dir: str) -> dict[str, str]:
    """Đọc tất cả file .docx trong thư mục input_dir.
    Trả về dict { filename: extracted_text }
    """
    results = {}
    input_path = Path(input_dir)

    if not input_path.exists():
        return results

    for filepath in sorted(input_path.glob("*.docx")):
        if filepath.name.startswith("~$"):
            continue
        try:
            results[filepath.name] = _extract_text(filepath)
        except Exception as e:
            results[filepath.name] = f"[Lỗi đọc file: {e}]"

    return results


def read_all_input_files(input_dir: str) -> dict[str, str]:
    """Đọc tất cả file .docx và .pdf trong thư mục input_dir.
    Trả về dict { filename: extracted_text }.
    """
    results = {}
    input_path = Path(input_dir)

    if not input_path.exists():
        return results

    filepaths = []
    for ext in (".docx", ".pdf"):
        filepaths.extend(input_path.glob(f"*{ext}"))

    for filepath in sorted(filepaths):
        if filepath.name.startswith("~$"):
            continue
        try:
            results[filepath.name] = _extract_text(filepath)
        except Exception as e:
            results[filepath.name] = f"[Lỗi đọc file: {e}]"

    return results


def _extract_text(filepath: Path) -> str:
    if filepath.suffix.lower() == ".pdf":
        return _extract_text_from_pdf(filepath)

    doc = Document(str(filepath))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_text_from_pdf(filepath: Path) -> str:
    reader = PdfReader(str(filepath))
    parts = []

    for page in reader.pages:
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            parts.append(text)

    return "\n".join(parts)


def read_notes(notes_path: str) -> str:
    """Đọc file ghi chú.md"""
    p = Path(notes_path)
    if p.exists():
        return p.read_text(encoding="utf-8")
    return ""
